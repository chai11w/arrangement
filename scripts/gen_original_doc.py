"""原封不动导出原始需求文档，不做任何修改"""
from docx import Document
from docx.shared import Pt

doc = Document()

style = doc.styles['Normal']
font = style.font
font.name = '宋体'
font.size = Pt(11)

with open('F:/cc/8-笔试/docs/arrangements-requirements.md', 'r', encoding='utf-8') as f:
    content = f.read()

for line in content.split('\n'):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.5
    run = p.add_run(line)
    run.font.name = '宋体'
    run.font.size = Pt(11)

output_path = 'F:/cc/8-笔试/即我安排模块-原始需求文档（未改动）.docx'
doc.save(output_path)
print(f'saved: {output_path}')
