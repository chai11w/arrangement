"""生成需求整理Word文档，重要内容标红"""
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

doc = Document()

# 设置默认字体
style = doc.styles['Normal']
font = style.font
font.name = '微软雅黑'
font.size = Pt(11)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

def add_heading_red(doc, text, level=1):
    """添加标题"""
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = '微软雅黑'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    return h

def add_para(doc, text, bold=False):
    """添加普通段落"""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.5
    run = p.add_run(text)
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    run.font.size = Pt(11)
    run.bold = bold
    return p

def add_red_text(doc, label, content):
    """添加带红色标签的段落：label普通，content红色加粗"""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.5
    if label:
        r1 = p.add_run(label)
        r1.font.name = '微软雅黑'
        r1._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
        r1.font.size = Pt(11)
    r2 = p.add_run(content)
    r2.font.name = '微软雅黑'
    r2._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    r2.font.size = Pt(11)
    r2.bold = True
    r2.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)
    return p

def add_mixed_para(doc, parts):
    """
    parts: list of (text, is_red, is_bold) tuples
    """
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.5
    for text, is_red, is_bold in parts:
        run = p.add_run(text)
        run.font.name = '微软雅黑'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
        run.font.size = Pt(11)
        if is_red:
            run.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)
        if is_bold:
            run.bold = True
    return p

# ========== 文档内容 ==========

add_heading_red(doc, '即我「安排」模块 — 原始需求整理文档', level=0)

# 一、项目背景
add_heading_red(doc, '一、项目背景', level=1)
add_mixed_para(doc, [
    ('本项目为即我（Arkme）App 的移动端 Demo，候选人需基于此 Demo，', False, False),
    ('在 Codex 客户端中完成「安排」模块的需求迭代', True, True),
    ('。', False, False),
])
add_para(doc, '当前项目仅保留首页、侧边栏、快记、洞见占位、我的等基础结构，「安排」模块完全空白，需要从零搭建。')

# 二、核心需求
add_heading_red(doc, '二、核心需求', level=1)

# 2.1 安排的定义
add_heading_red(doc, '2.1「安排」的定义', level=2)
add_mixed_para(doc, [
    ('核心洞察：', False, True),
    ('待办、日程、项目、任务、提醒、规划等传统概念，本质上都是"还没发生、但需要用户后续执行落地的大大小小的东西"。因此，用「安排」这个更抽象、更统一的概念来承载所有这些。', True, True),
])

# 2.2 AI 识别安排
add_heading_red(doc, '2.2 AI 自动识别安排 ⭐', level=2)
add_red_text(doc, '【核心能力】', 'AI 从发给自己的消息、私聊、群聊等对话中，自动识别出安排内容并生成安排项。')
add_para(doc, '场景示例：')
add_para(doc, '① 发给自己："后天去一趟医院" → 自动生成相关日程')
add_mixed_para(doc, [
    ('② 私聊：对方说"明天来公司帮我带个早餐"，自己回复"好的" → ', False, False),
    ('自己的安排模块生成"明天到公司帮xxx带早餐"；对方安排模块生成"yy明天来公司帮自己带早餐"', True, True),
])
add_mixed_para(doc, [
    ('③ 群聊：类似识别逻辑，', False, False),
    ('需判断是只展示与自己相关的安排，还是展示群聊中所有安排', True, True),
    ('。', False, False),
])
add_mixed_para(doc, [
    ('④ 多人对话连续识别：', False, True),
    ('如连续对话中说"帮对象带A、B、C、D、E几个物品"，需要识别出完整的物品清单。这是考察对AI理解和多轮对话处理的关键场景。', True, True),
])

# 2.3 归集与合并
add_heading_red(doc, '2.3 归集与合并机制 ⭐', level=2)
add_red_text(doc, '【核心机制】', '多条相关的安排应合并为同一条，但在详情中保留所有关联对话上下文，让用户感知到对应了哪些内容。')
add_para(doc, '场景示例：')
add_para(doc, '用户规划了"后天去医院"→ 爸爸发消息"一定记得去医院"→ 自己回复"好的，有安排了"')
add_para(doc, '→ 姐姐问"身体情况怎么办"→ 自己回复"已经挂号了"')
add_mixed_para(doc, [
    ('→ 这三条应合并为同一条安排，', False, False),
    ('详情中完整呈现所有关联对话', True, True),
    ('，让用户知道哪些对话触发了这个安排。', False, False),
])

# 2.4 手动创建
add_heading_red(doc, '2.4 手动创建机制 ⭐', level=2)
add_red_text(doc, '【必须保留】', 'AI 存在局限性，无法识别隐喻性输入（如用户用"～～"代表"游泳"），因此必须保留手动创建安排的能力。')
add_para(doc, '手动创建是基础功能的底线保障。')

# 2.5 时间与日历
add_heading_red(doc, '2.5 时间识别与日历呈现 ⭐', level=2)
add_red_text(doc, '【关键交互】', '每个安排项应尽可能识别出时间，或让用户后续输入时间。需要日历视图，让用户有全局总览感。')
add_para(doc, '时间类型区分：')
add_para(doc, '• 待办/任务 → 截止时间形式')
add_para(doc, '• 日程 → 时间段形式')
add_para(doc, '• 提醒 → 触发时间点的机制（非完成导向）')

# 2.6 完成机制
add_heading_red(doc, '2.6 完成机制 ⭐', level=2)
add_mixed_para(doc, [
    ('核心原则：', False, True),
    ('不能永无止境创造新安排导致堆积。', True, True),
])
add_para(doc, '两种完成方式：')
add_para(doc, '① 手动完成：用户进入模块手动操作，需设计优雅的交互方式')
add_mixed_para(doc, [
    ('② AI 自动判断完成：', False, False),
    ('根据对话上下文自动识别。如用户跟姐姐说"今天上午去医院体检了，身体没啥情况" → 自动将"去医院检查"安排项标记为完成。', True, True),
])

# 2.7 提醒机制
add_heading_red(doc, '2.7 提醒机制', level=2)
add_para(doc, '提醒是一种中间态机制，与待办/任务的完成导向不同：')
add_para(doc, '• 用户可明确创建纯粹提醒（到时间触发即可）')
add_para(doc, '• 任务/待办/日程也可以设置提醒（提前/循环提醒），让用户记住去执行')

# 2.8 AI 执行分级
add_heading_red(doc, '2.8 AI 执行分级', level=2)
add_red_text(doc, '【进阶能力】', '安排项需区分三类：')
add_para(doc, '① 只能靠用户自己人工完成')
add_para(doc, '② AI 可先帮用户执行部分，用户完成后续')
add_para(doc, '③ AI 可百分百替用户直接执行')
add_para(doc, '不同类别需有不同的处理机制。')

# 2.9 减轻用户焦虑
add_heading_red(doc, '2.9 减轻用户焦虑 — 产品哲学 ⭐⭐', level=2)
add_red_text(doc, '【最重要设计原则】', '这是贯穿整个模块的核心产品哲学，需在所有交互设计中贯彻。')
add_para(doc, '具体要点：')
add_mixed_para(doc, [
    ('① ', False, False),
    ('不采用红色逾期提醒：', True, True),
    ('传统任务管理软件中，80%的任务不能在截止时间完成，大量红色逾期提醒让用户产生焦虑感，想远离App。', False, False),
])
add_mixed_para(doc, [
    ('② ', False, False),
    ('"以后再说"机制：', True, True),
    ('用户看到超期未完成项，既不能点完成、又不能删除、又不能快速想到新截止时间时，内心最直觉的反应是"以后再说吧"。应提供长按/右滑触发"以后再说"的优雅操作。', False, False),
])
add_mixed_para(doc, [
    ('③ ', False, False),
    ('无轻重缓急罗列：', True, True),
    ('不强制所有安排平铺列出，理论上真正值得关注的核心安排是少数。', False, False),
])
add_mixed_para(doc, [
    ('④ ', False, False),
    ('人生哲学：', True, True),
    ('人生充满变数，不一定必须遵循之前设定的轨迹。不完成又能怎样？产品层面不应约束用户。', False, False),
])

# 2.10 API 绑定
add_heading_red(doc, '2.10 API 绑定 ⭐', level=2)
add_red_text(doc, '【第二层核心能力】', '每个用户需要能够输入自己的大模型 API Key，消耗自己的 Token 来驱动 AI 识别功能。')
add_para(doc, '这是考察候选人对 AI 理解和使用的关键环节。')

# 三、实现策略
add_heading_red(doc, '三、实现策略', level=1)

add_heading_red(doc, '3.1 版本迭代节奏 ⭐', level=2)
add_red_text(doc, '【最小可用原则，小步快跑】', '不要试图一次实现所有需求，而要逐步迭代。')
add_para(doc, '建议分层：')
add_mixed_para(doc, [
    ('第一层（框架 + 基础交互）：', False, True),
    ('搭建安排模块框架 → 手动创建安排 → 手动完成安排。', True, True),
    ('这一步用于判断候选人的交互界面审美和基本功。', False, False),
])
add_mixed_para(doc, [
    ('第二层（AI 集成）：', False, True),
    ('增加 API 绑定功能 → 逐个场景实现 AI 识别。', True, True),
    ('一个场景一个场景做，不要贪多，每做一个场景对 AI 的理解就提升一点。', False, False),
])
add_mixed_para(doc, [
    ('第三层（体验打磨）：', False, True),
    ('完善合并、日历、提醒、完成判断、焦虑处理等进阶功能。', True, True),
])

add_heading_red(doc, '3.2 质量要求', level=2)
add_para(doc, '• 视觉与交互：尽可能简约、优雅、舒服，注重用户体验细节')
add_para(doc, '• 宁可做少但要打磨好，而非做全但难用')
add_para(doc, '• 把自己当成用户，站在用户角度换位思考')
add_para(doc, '• 市面上无可借鉴方案，需独立思考')
add_para(doc, '• AI功能需自测能跑通')

# 四、评估重点
add_heading_red(doc, '四、评估重点', level=1)
add_para(doc, '1. 对需求的理解深度（第一层理解了多少）')
add_para(doc, '2. 与 AI 探讨理解需求后的执行能力')
add_para(doc, '3. 交互的优雅程度')
add_para(doc, '4. 视觉层面的美观舒服')
add_para(doc, '5. 用户体验层面的好用程度')
add_para(doc, '6. 换位思考能力 — 将自己放在用户角色')

# 五、关键信息速查
add_heading_red(doc, '五、关键信息速查', level=1)
add_para(doc, '• GitHub 地址：https://github.com/imanai666/ArkmeDemo')
add_para(doc, '• 本地测试入口：http://127.0.0.1:5173/ （移动端Demo）')
add_para(doc, '• 消息测试后台：http://127.0.0.1:5173/sendtest')
add_para(doc, '• 技术栈：React 18 + Vite + Tailwind CSS + TypeScript')
add_para(doc, '• 包管理：pnpm（>=9.12.3）')
add_para(doc, '• Node：>=20.10.0')

# 保存
output_path = 'F:/cc/8-笔试/即我安排模块-需求整理文档.docx'
doc.save(output_path)
print(f'文档已保存至: {output_path}')
